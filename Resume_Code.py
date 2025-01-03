from docx import Document
from tkinter import Tk
from tkinter.filedialog import asksaveasfilename

# Create a new Word document
doc = Document()

# Add content to the Word document
doc.add_heading('Jatin Chandani', level=0)
doc.add_paragraph('Colchester, UK | jatinchandani8@gmail.com | +44 7407 022519 | LinkedIn Profile: https://www.linkedin.com/in/jatinchandani28')
doc.add_paragraph('---')

doc.add_heading('Curriculum Vitae', level=1)

doc.add_heading('Professional Summary', level=2)
doc.add_paragraph("Dynamic Full-Stack Developer and Data Consultant with over 4 years of experience in data analysis, project management, and software development. "
                  "Proficient in leveraging analytical skills and technology to deliver impactful solutions. Skilled in interpreting and synthesizing complex data to "
                  "support evidence-based decision-making, with a growing interest in biodiversity conservation and environmental policy. Experienced in working with "
                  "international clients across diverse cultural contexts, demonstrating strong communication and team collaboration skills.")

doc.add_paragraph('---')
doc.add_heading('Key Skills', level=2)

doc.add_heading('Analytical and Policy-Related Skills', level=3)
doc.add_paragraph('- Data collection, analysis, and synthesis for evidence-informed decision-making\n'
                  '- Research and report writing for technical and non-technical audiences\n'
                  '- Experience in creating dashboards and visualizations for strategic insights')

doc.add_heading('Technical Skills', level=3)
doc.add_paragraph('- Programming: Python, R, JavaScript, TypeScript\n'
                  '- Data Analysis: Pandas, NumPy, Matplotlib, Power BI, SQL\n'
                  '- Cloud Computing: AWS (S3, RDS, EC2, Lambda, CloudFormation)\n'
                  '- Tools: NVivo, Tableau, RESTful APIs, Excel (Advanced)')

doc.add_heading('Soft Skills', level=3)
doc.add_paragraph('- Strong attention to detail and time management\n'
                  '- Cross-functional collaboration and stakeholder engagement\n'
                  '- Excellent written and verbal communication')

doc.add_paragraph('---')
doc.add_heading('Professional Experience', level=2)

doc.add_heading('Freelance Full-Stack Developer & Data Consultant', level=3)
doc.add_paragraph('Mar 2023 – Present | UK, Dubai, India')
doc.add_paragraph('- Designed and deployed AWS-based cloud architectures to streamline workflows and optimize efficiency, achieving a 20% reduction in costs.\n'
                  '- Synthesized data from various sources to create actionable insights for decision-makers in cross-functional projects.\n'
                  '- Developed serverless applications with AWS Lambda, enhancing scalability and response times.\n'
                  '- Delivered presentations and technical reports to clients, ensuring clarity and alignment with project goals.')

doc.add_heading('Software Developer', level=3)
doc.add_paragraph('Engineer Master Solutions Pvt. Ltd. | Jul 2022 – Feb 2023')
doc.add_paragraph('- Led data-driven initiatives by creating dashboards and performing data analysis to improve decision-making processes.\n'
                  '- Enhanced backend performance by optimizing database structures, reducing latency by 35%.\n'
                  '- Developed and implemented CI/CD pipelines, ensuring faster deployment cycles.')

doc.add_heading('Full-Stack Developer Intern', level=3)
doc.add_paragraph('Engineer Master Solutions Pvt. Ltd. | Mar 2020 – Mar 2021')
doc.add_paragraph('- Contributed to the development of MERN stack applications with a focus on responsive design and data-driven solutions.\n'
                  '- Analyzed and structured data to support backend system improvements, reducing downtime by 15%.')

doc.add_paragraph('---')
doc.add_heading('Key Projects', level=2)

doc.add_heading('Data-Driven Biodiversity Dashboard (Personal Project)', level=3)
doc.add_paragraph('- Developed an interactive dashboard using Python and Power BI to visualize biodiversity indicators for policymaking.\n'
                  '- Integrated publicly available biodiversity data sets to highlight trends and support evidence-based conservation strategies.')

doc.add_heading('AI-Driven Real Estate Platform', level=3)
doc.add_paragraph('- Utilized AWS SageMaker for predictive analytics, delivering higher lead conversions and improved recommendations.\n'
                  '- Demonstrated the ability to analyze large datasets and deliver actionable insights.')

doc.add_heading('Cloud Migration Projects', level=3)
doc.add_paragraph('- Migrated enterprise systems to AWS using CloudFormation templates, enhancing scalability and deployment efficiency.')

doc.add_paragraph('---')
doc.add_heading('Education', level=2)

doc.add_heading('MSc in Applied Data Science', level=3)
doc.add_paragraph('University of Essex, UK | October 2023 – October 2024')
doc.add_paragraph('- Coursework in data visualization, machine learning, and data-driven decision-making.\n'
                  '- Focused on data analysis for real-world challenges.')

doc.add_heading('Bachelor of Business Administration (BBA)', level=3)
doc.add_paragraph('Acropolis Institute of Management Studies & Research, India | July 2016 – May 2019')
doc.add_paragraph('- Specialization in Finance and Marketing, emphasizing business management and strategic decision-making.\n'
                  '- Gained practical experience through case studies and projects focused on market analysis and business operations.\n'
                  '- Developed leadership skills by participating in and organizing academic and extracurricular events.\n'
                  '- Awards and Achievements:\n'
                  '  - Winner of IRIS 2018 (IIM Indore)\n'
                  '  - Winner of Prodigy (IIM Bangalore).\n'
                  '- Participated in national-level business plan competitions, showcasing innovative business strategies.')

doc.add_paragraph('---')
doc.add_heading('Certifications and Training', level=2)
doc.add_paragraph('- AWS Certified Solutions Architect – Associate\n'
                  '- IBM Data Science Professional Certificate\n'
                  '- Machine Learning Foundations (LinkedIn Learning)\n'
                  '- Ongoing self-study in biodiversity policy and environmental governance.')

doc.add_paragraph('---')
doc.add_heading('Achievements', level=2)
doc.add_paragraph('- Improved application scalability and reduced costs by leveraging serverless technologies.\n'
                  '- Delivered technical reports and presentations tailored to diverse audiences, fostering collaboration and understanding.\n'
                  '- Successfully led cross-functional teams to meet critical deadlines in fast-paced environments.')

doc.add_paragraph('---')
doc.add_heading('Interests and Volunteering', level=2)
doc.add_paragraph('- Actively exploring biodiversity and environmental policies through online courses and self-initiated research.\n'
                  '- Volunteered in organizing community workshops focusing on sustainability and digital literacy.')

doc.add_paragraph('---')
doc.add_heading('References', level=2)
doc.add_paragraph('Available upon request.')

# Use tkinter to open a save file dialog
Tk().withdraw()  # Hide the main tkinter window
file_path = asksaveasfilename(
    defaultextension=".docx",
    filetypes=[("Word Documents", "*.docx")],
    title="Save Your CV",
    initialfile="Jatin_Chandani_Updated_CV.docx"
)

if file_path:
    # Save the document to the selected path
    doc.save(file_path)
    print(f"File saved at: {file_path}")
else:
    print("File save operation canceled.")